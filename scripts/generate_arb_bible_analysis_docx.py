"""Generate Polymarket Arbitrage Bible analysis Word document for Grok-Bot-1."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT = Path(__file__).resolve().parents[1] / "Polymarket_Arbitrage_Bible_Analysis.docx"


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "D5E8F0") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Polymarket Arbitrage Bible\nDeep-Dive Analysis")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Grok-Bot-1 / Hermes BTC Pulse Engine\n").font.size = Pt(12)
    sub.add_run(f"Prepared: {date.today().isoformat()}\n").font.size = Pt(11)
    sub.add_run("Source: Polymarket Arbitrage Bible.pdf (Roan / arXiv 2508.03474)").font.size = Pt(10)

    doc.add_page_break()

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "The Polymarket Arbitrage Bible compiles Roan's roadmap and the research paper "
        "\"Unravelling the Probabilistic Forest: Arbitrage in Prediction Markets\" (arXiv:2508.03474). "
        "The central thesis is that profitable Polymarket arbitrage is not a simple YES+NO addition problem. "
        "Production systems require combinatorial constraint modeling (integer programming over the marginal polytope), "
        "Bregman/KL projection for optimal trade sizing, Frank-Wolfe iteration with an IP oracle (e.g. Gurobi), "
        "and execution realism (VWAP, non-atomic CLOB fills, depth caps, modified Kelly)."
    )
    doc.add_paragraph(
        "Empirical evidence from April 2024–April 2025 shows roughly $40M extracted across ~17,218 conditions, "
        "with 41% exhibiting single-market mispricing and the top bot earning ~$2M from ~4,000 trades. "
        "The mathematical infrastructure—not raw speed alone—defines the competitive gap."
    )
    doc.add_paragraph(
        "Grok-Bot-1's pulse engine implements a deliberately scoped subset: within-window dutch-book arbitrage "
        "on BTC 5m/15m up/down markets (vwap_up + vwap_down < $1), with a segregated paper ArbLedger. "
        "This aligns with Roan's note that for 2-outcome windows the marginal polytope is trivial. "
        "The Bible's full combinatorial stack remains out of scope but defines a clear multi-tier roadmap "
        "if the bot expands beyond short-horizon BTC windows."
    )

    # Chapter summaries
    doc.add_heading("Chapter 1: The Marginal Polytope Problem", level=1)
    doc.add_paragraph(
        "Single-market efficiency (YES + NO = $1) is necessary but insufficient. Cross-market logical "
        "dependencies—subset, implication, mutual exclusion—create pricing constraints that span multiple "
        "conditions. The space of valid price vectors is the marginal polytope M; prices outside M are exploitable."
    )
    add_bullets(doc, [
        "Example: \"Trump wins PA\" vs \"GOP wins PA by 5+ pts\" — landslide YES is a subset of win YES, "
        "so P(landslide) cannot exceed P(win). Violations enable risk-free trades.",
        "Brute enumeration of 2^n outcomes is infeasible (2^63 for NCAA 2010; 5,000+ live markets).",
        "Integer programming encodes validity with linear constraints (e.g. exactly-one-outcome, "
        "semifinal exclusion) replacing exponential search.",
        "Empirical: 17,218 conditions scanned; 7,051 (41%) had single-market arbitrage; "
        "median deviation $0.60 vs $1.00 nominal; 13 confirmed cross-market pairs.",
    ])

    doc.add_heading("Chapter 2: Bregman Projection", level=1)
    doc.add_paragraph(
        "Detecting mispricing is separate from computing the optimal arbitrage trade. Euclidean distance "
        "treats all cent moves equally; LMSR-based markets require Bregman divergence (KL for log utilities), "
        "weighting moves near 0 and 1 more heavily."
    )
    add_bullets(doc, [
        "Core result: maximum guaranteed profit equals Bregman distance from current prices to the "
        "arbitrage-free space.",
        "Projection yields direction (buy/sell which legs), size (depth-aware), and profit bound.",
        "Top arbitrator: ~$2,009,632 in one year solving this optimization faster than competitors.",
    ])

    doc.add_heading("Chapter 3: Frank-Wolfe Algorithm", level=1)
    doc.add_paragraph(
        "Direct Bregman projection onto M is intractable because M has exponentially many vertices. "
        "Frank-Wolfe iteratively grows a small active set of valid outcomes, querying an IP solver each step."
    )
    add_bullets(doc, [
        "Each iteration adds one vertex; 50–150 iterations typically suffice.",
        "Gurobi solve times: <1s early, 10–30s mid-tournament, <5s late as feasible set shrinks.",
        "Barrier Frank-Wolfe handles LMSR gradient explosion near price boundaries (adaptive ε shrinkage).",
        "FWMM outperformed LCMM by 38% once IP solves fit inside the trading window (~30 min NCAA projection).",
    ])

    doc.add_heading("Chapter 4: Execution Reality", level=1)
    doc.add_paragraph(
        "Even correct mathematical arbitrage can lose money at execution. Polymarket's CLOB is non-atomic; "
        "legs fill sequentially and move the book."
    )
    add_bullets(doc, [
        "Minimum edge threshold ~$0.05 to survive execution risk (paper filter).",
        "VWAP over ask/bid ladders is the executable price, not top-of-book quotes.",
        "Profit capped by thinnest leg depth; cross-market arb needs simultaneous liquidity on all positions.",
        "Position sizing: modified Kelly with execution probability p; cap at ~50% of book depth.",
    ])

    doc.add_heading("Chapter 5: Production System Architecture", level=1)
    doc.add_paragraph("The deployed research stack has three layers plus infrastructure:")
    add_numbered(doc, [
        "Data pipeline: WebSocket order books + Polygon historical events (86M trades analyzed).",
        "Dependency detection: LLM screening (DeepSeek-R1-Distill-Qwen-32B, ~81.45% accuracy) "
        "with manual verification — 1,576 dependent pairs from 46,360 candidates.",
        "Layer 1 — LCMM: millisecond linear constraint checks (sum-to-one, implication).",
        "Layer 2 — Frank-Wolfe + Gurobi: α=0.9 extraction, ε=0.1 initial barrier, 1e-6 convergence.",
        "Layer 3 — Execution validation: simulate fills, slippage, min $0.05 profit gate.",
    ])

    # Empirical table
    doc.add_heading("Empirical Results (Apr 2024 – Apr 2025)", level=1)
    add_table(doc,
              ["Category", "Extracted Profit (USD)", "Notes"],
              [
                  ["Single-condition (buy both low)", "$5,899,287", "Dutch-book style"],
                  ["Single-condition (sell both high)", "$4,682,075", "Mint set, sell bids"],
                  ["Market rebalancing (YES/NO baskets)", "$29,011,589", "Cross-leg rebalancing"],
                  ["Cross-market combination", "$95,634", "13 verified pairs"],
                  ["Total", "$39,688,585", "~$40M headline"],
                  ["Top 10 bots share", "$8,127,849", "20.5% of total"],
                  ["Top single bot", "$2,009,632", "4,049 trades, ~$496 avg"],
              ])

    add_table(doc,
              ["Market Statistics", "Value"],
              [
                  ["Conditions analyzed", "17,218"],
                  ["Single-market arb conditions", "7,051 (41%)"],
                  ["Median pricing deviation", "$0.60 (vs $1.00)"],
                  ["Cross-market pairs screened", "1,576 dependent"],
                  ["Manually verified exploitable pairs", "13"],
                  ["LLM dependency screening accuracy", "81.45%"],
              ])

    # Glossary
    doc.add_heading("Concept Glossary", level=1)
    add_table(doc,
              ["Term", "Definition"],
              [
                  ["Marginal polytope", "Convex hull of valid price vectors consistent with all logical outcome constraints."],
                  ["Integer programming", "Linear constraints over binary outcome indicators; avoids 2^n enumeration."],
                  ["Bregman / KL divergence", "Information-theoretic distance for probability prices; LMSR-consistent."],
                  ["LMSR", "Logarithmic Market Scoring Rule — Polymarket MM pricing backbone."],
                  ["Frank-Wolfe", "Iterative convex optimizer adding one polytope vertex per step via IP oracle."],
                  ["LCMM", "Linear Constraint Market Maker — fast first-pass mispricing filter."],
                  ["FWMM", "Frank-Wolfe Market Maker — full Bregman projection engine."],
                  ["CLOB", "Central limit order book; sequential, non-atomic fills."],
                  ["VWAP", "Volume-weighted average price across ladder depth."],
                  ["Dutch book", "Buy all mutually exclusive outcomes for < $1 guaranteed payout."],
              ])

    # Grok-Bot-1 gap analysis
    doc.add_heading("Grok-Bot-1 Gap Analysis", level=1)
    doc.add_paragraph(
        "The Hermes BTC Pulse engine (plugins/hermes-trading-engine/engine/pulse/) runs two segregated "
        "strategies: directional (TradingView + edge model) and within-window arbitrage. "
        "Arbitrage is PAPER ONLY with deterministic P&L in a separate ArbLedger."
    )

    doc.add_heading("What Grok-Bot-1 Implements Today", level=2)
    add_bullets(doc, [
        "detect_arbitrage(): buy-both (asks) and sell-both (mint $1 set, sell bids) on up/down books.",
        "VWAP fill simulation via execution_gate.vwap_fill — not top-of-book fantasy.",
        "Depth cap: max_depth_consume_frac=0.5 of thinner leg; full-fill required for actionable.",
        "Stale-book guard (max_book_age_s), tick alignment, epsilon edge filter (default arb_epsilon=0.01).",
        "Separate ArbLedger: never blended into directional WR/PF stats; settle at window close.",
        "Engine integration: arb preempts directional on same window; stop_conditions per strategy.",
        "Scope: BTC 5m/15m INDEX:BTCUSD windows only — 2-outcome marginal polytope is trivial.",
    ])

    doc.add_heading("Alignment with the Bible", level=2)
    add_table(doc,
              ["Bible Concept", "Grok-Bot-1 Status"],
              [
                  ["Single-market dutch book (YES+NO ≠ $1)", "Implemented (VWAP, depth, epsilon)"],
                  ["VWAP execution pricing", "Implemented (shared with execution gate)"],
                  ["Non-atomic / partial fill awareness", "Partial — requires full dual-leg fill; no live order placement"],
                  ["Depth-limited sizing (~50% book)", "Implemented (max_depth_consume_frac=0.5)"],
                  ["Min profit threshold ($0.05 paper)", "Configurable (arb_epsilon; live default 0.01)"],
                  ["Marginal polytope / cross-market deps", "Not implemented — out of scope for 2-outcome BTC"],
                  ["Integer programming constraints", "Not implemented"],
                  ["Bregman / KL projection", "Not implemented"],
                  ["Frank-Wolfe + Gurobi", "Not implemented"],
                  ["LLM dependency screening", "Not implemented (Grok used for directional, not arb deps)"],
                  ["Modified Kelly sizing", "Not implemented for arb (fixed size_usd / max_usd caps)"],
                  ["Multi-condition scanning (17k+)", "Not implemented (2 markets: 5m + 15m BTC)"],
                  ["Live execution", "PAPER ONLY — simulates fills, no CLOB order submission"],
              ])

    doc.add_heading("Strategic Assessment", level=2)
    doc.add_paragraph(
        "For short-horizon BTC up/down windows, Roan's own framing acknowledges that combinatorial "
        "machinery is unnecessary: arbitrage reduces to vwap_up + vwap_down < 1. Grok-Bot-1 correctly "
        "implements this tier with execution realism borrowed from the directional gate. The competitive "
        "gap versus the $40M extraction study is therefore expected: BTC 5m/15m windows are a tiny, "
        "highly competitive subset with shallow books and sub-second arb decay."
    )
    doc.add_paragraph(
        "The Bible's infrastructure becomes relevant if Grok-Bot-1 expands to multi-condition event clusters "
        "(elections, sports brackets, correlated macro markets) or deploys live CLOB execution with "
        "parallel leg submission and inventory across windows."
    )

    # Roadmap
    doc.add_heading("Optional Roadmap Tiers", level=1)
    add_table(doc,
              ["Tier", "Scope", "Effort", "Value for BTC Pulse"],
              [
                  ["T0 (current)", "2-outcome dutch book, paper, VWAP", "Done", "High — correct scope for 5m/15m"],
                  ["T1", "Live CLOB execution, parallel legs, $0.05 min edge", "Medium", "Medium — needs latency + capital"],
                  ["T2", "Cross-window / cross-market BTC correlation", "High", "Low — few logical deps on BTC windows"],
                  ["T3", "LCMM layer for N correlated Polymarket events", "High", "High if expanding beyond BTC"],
                  ["T4", "Frank-Wolfe + Gurobi Bregman projection", "Very high", "High for multi-condition only"],
                  ["T5", "LLM dependency screening pipeline", "Medium", "Prerequisite for T3/T4"],
              ])

    doc.add_heading("Recommended Near-Term Actions", level=2)
    add_numbered(doc, [
        "Keep T0 scope locked for BTC Pulse; do not port Gurobi into the hot loop without multi-condition markets.",
        "Raise arb_epsilon toward 0.05 before any live execution to match paper's execution-risk filter.",
        "Instrument arb scan telemetry (scans, near_miss, min_vwap_residual) — already in ArbLedger; monitor on soak.",
        "If expanding markets: add T5 dependency screening before T3/T4; manual verify the 13-pair workflow.",
        "Segregate arb P&L reporting remains correct — preserve separation from directional scoring.",
    ])

    # References
    doc.add_heading("References", level=1)
    refs = [
        ("Original thread (Roan)", "https://x.com/RohOnChain/status/2017314080395296995"),
        ("Research paper", "https://arxiv.org/abs/2508.03474"),
        ("IP market making foundation", "https://arxiv.org/abs/1606.02825"),
        ("LMSR explainer", "https://www.cultivatelabs.com/crowdsourced-forecasting-guide/how-does-logarithmic-market-scoring-rule-lmsr-work"),
        ("Bregman divergences", "https://mark.reid.name/blog/meet-the-bregman-divergences.html"),
        ("KL divergence", "https://en.wikipedia.org/wiki/Kullback%E2%80%93Leibler_divergence"),
        ("Frank-Wolfe algorithm", "https://en.wikipedia.org/wiki/Frank%E2%80%93Wolfe_algorithm"),
        ("Gurobi optimizer", "https://www.gurobi.com/"),
        ("Polymarket CLOB docs", "https://docs.polymarket.com/"),
        ("VWAP", "https://www.investopedia.com/terms/v/vwap.asp"),
        ("Kelly criterion", "https://www.investopedia.com/articles/trading/04/091504.asp"),
        ("Decrypt $40M report", "https://decrypt.co/339958/40-million-free-money-glitch-crypto-prediction-markets"),
        ("Grok-Bot-1 arbitrage.py", "hermes-agent-main/plugins/hermes-trading-engine/engine/pulse/arbitrage.py"),
    ]
    for label, url in refs:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{label}: ").bold = True
        p.add_run(url)

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("— End of analysis —").italic = True

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()